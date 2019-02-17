from django.shortcuts import render
from gtts import gTTS
import os
from pydub import AudioSegment
import threading
import sys
import PyPDF2
import time
from pdf2image import convert_from_path
from PIL import Image
from pytesseract import image_to_string
import multiprocessing
from docx import Document
import re
try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile
from flask import Flask
from flask import render_template


WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'


def generate_single_track():
    if os.path.isfile("audio/piece0.mp3"):
        # It will start reading the file aloud right when it makes the first one
        # and then it will save the full export when its done.
        i = 1
        # Initalizeing the full track
        full_track = AudioSegment.from_mp3("audio/piece0.mp3")
        while os.path.isfile("audio/piece" + str(i) + ".mp3"):
            # reads in a new track
            new_track = AudioSegment.from_mp3("audio/piece" + str(i) + ".mp3")
            print("adding track num " + str(i))
            # adds the new tack to the full exported thing
            full_track += new_track
            i += 1
        folder, filename = os.path.split(sys.argv[1])
        full_track.export("audio/" + filename + ".mp3", format='mp3')


def convert_txt_to_docx():
    path = '~/github/Text-to-Speech/output.txt'
    direct = os.listdir(path)

    for i in direct:
        document = Document()
        document.add_heading(i, 0)
        myfile = open("output.txt").read()
        myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
        # p = document.add_paragraph(myfile)
        document.save('output.docx')


def delete_old_files():
    # Remove any leftover audio
    i = 0
    while os.path.isfile("/home/c/github/Text-to-Speech/audio/piece" + str(i) + ".mp3"):
        os.remove("/home/c/github/Text-to-Speech/audio/piece" + str(i) + ".mp3")
        i += 1

    # Remove any leftover images
    i = 0
    while os.path.isfile("/home/c/github/Text-to-Speech/images/test" + str(i) + ".jpg"):
        os.remove("/home/c/github/Text-to-Speech/images/test" + str(i) + ".jpg")
        i += 1


def get_docx_text(path):
    """
    Take the path of a docx file as argument, return the text in unicode.
    """
    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    tree = XML(xml_content)

    paragraphs = []
    for paragraph in tree.getiterator(PARA):
        texts = [node.text
                 for node in paragraph.getiterator(TEXT)
                 if node.text]
        if texts:
            paragraphs.append(''.join(texts))

    return '\n\n'.join(paragraphs)


# Process the audio
def ask_google(string, i):
    tts = gTTS(text=string, lang='en')
    tts.save("/home/c/github/Text-to-Speech/audio/piece" + str(i) + ".mp3")


# Save the image and then turn it to text
def image_pdf(image, i):
    image.save("/home/c/github/Text-to-Speech/images/test" + str(i) + ".jpg")
    image = Image.open("/home/c/github/Text-to-Speech/images/test" + str(i) + ".jpg", mode='r')
    # print(image_to_string(image))
    return image_to_string(image)


# Turn the large body of text into small pieces
def make_phrases(s):
    word_list = s.split(" ")
    phrase = ""
    phrase_list = []
    for i in range(len(word_list)):
        phrase += word_list[i] + " "
        if len(phrase) >= 110:
            phrase_list.append(phrase)
            phrase = ""
    return phrase_list


# Multithread it
def make_threads(phrase_list, threadingCounter):
    threadingCounterDefault = threadingCounter
    length = len(phrase_list)
    i = 0
    cores = multiprocessing.cpu_count()
    print (length)
    while i != length:
        threads = threading.activeCount()
        # print ("THe threads are " + threads + "\n" + "The cores are " + cores)
        print (str(threadingCounter + 1) + "/" + str(length + threadingCounterDefault))
        t1 = threading.Thread(target=ask_google, args=(phrase_list[i], threadingCounter,))
        t1.start()
        print("There are currently " + str(threads) + " threads running")
        threadingCounter += 1
        i += 1
    try:
        t1.join()
    except UnboundLocalError:
        print("Tried to allow threads to close when none existed")

    return threadingCounter - threadingCounterDefault

def main(filename):
    print("Running the main loop")
    s = ""
    f = open("/home/c/github/Text-to-Speech/output.txt", 'w')
    try:
        if (filename[-4:] == ".pdf"):
            pdfReader = PyPDF2.PdfFileReader(open(filename, 'rb'))
            threadingCounter = 0
            x = 0
            i = 0
            for i in range(5):
                y = len(str(pdfReader.getPage(5 + i * 2).extractText()))
                if x < y:
                    x = y;
            if x > 65:
                '''Extracts the text from the pdf, splits it into small enough
                pieces for it to go to google and then multithreads the sending of
                the files to google and saves the replys from google to mp3s'''
                print("This is a normal pdf.")
                for i in range(pdfReader.numPages):
                    # Better for debuging
                    s = str(pdfReader.getPage(i).extractText())
                    ph = make_phrases(s)
                    threadingCounter += make_threads(ph, threadingCounter)
                    f.write(s)
            else:
                ''' Turns the images into text and them into small enough sizes
                to send to google and then multi threads the process of sending
                the strings to google than saves those files to mp3s'''
                print("This is a scanned in pdf.")
                images = convert_from_path(filename)
                for i in range(len(images)):
                    text = (image_pdf(images[i], i))
                    f.write(text)
                    print(text)
                    ph = make_phrases(text)
                    threadingCounter += make_threads(ph, threadingCounter)


        elif filename[-4:] == ".txt":
            # Better for debuging. Make into one line when done
            '''Converts the text into small enough pieces to send to google, multi
             threads the sending of those files to google and then saves those files
             google exports to mp3s'''
            f = open(filename, "r")
            ph = make_phrases(f.read())
            make_threads(ph, 0)
            # o = open("output.txt", 'w')
            o.write(f.read())

        elif filename[-5:] == ".docx":
    #       Better for debuging when done put it in one function call
            '''Extracts the text out of the docx file and then splits that into
            phrases small enough to send to google and then multithreads that
            process and then saves the exported files as mp3s'''
            text = get_docx_text(filename)
            ph = make_phrases(text)
            make_threads(ph, 0)
            f.write(text)
        i = 1
        # Initalizeing the full track
        full_track = AudioSegment.from_mp3("/home/c/github/Text-to-Speech/audio/piece0.mp3")
        while os.path.isfile("audio/piece" + str(i) + ".mp3"):
            # reads in a new track
            new_track = AudioSegment.from_mp3("/home/c/github/Text-to-Speech/audio/piece" + str(i) + ".mp3")
            print("adding track num " + str(i))
            # adds the new tack to the full exported thing
            full_track += new_track
            i += 1
        folder, filename = os.path.split(filename)
        full_track.export("/home/c/github/Text-to-Speech/audio/" + filename + ".mp3", format='mp3')
        # f = open("/home/c/github/Text-to-Speech/output1.txt", 'r')
        # o = open("/home/c/github/Text-to-Speech/output.txt", 'w')
        o.write(f.read())
        delete_old_files()

    except KeyboardInterrupt:
        print("Goodbye World\n")
        delete_old_files()

main(sys.argv[1])
